from __future__ import annotations

from pathlib import Path

from app.ingestion.parser_router import ParserRouter


def _text(chunks) -> str:
    return "\n".join(chunk.text for chunk in chunks)


def test_csv_parser_keeps_headers(tmp_path: Path) -> None:
    path = tmp_path / "parts.csv"
    path.write_text("object,part,article_number\nнасос NPK-200,корпус,ART-NPK-200-BODY\n", encoding="utf-8")
    parsed = ParserRouter().parse_document(str(path), doc_id="csv")
    text = _text(parsed.chunks)
    assert "Table columns:" in text
    assert "object: насос NPK-200" in text
    assert parsed.chunks[0].metadata["table_id"]
    assert parsed.chunks[0].metadata["row_id"] == "0"


def test_xlsx_parser_keeps_headers(tmp_path: Path) -> None:
    from openpyxl import Workbook

    path = tmp_path / "materials.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["material", "standard"])
    ws.append(["12Х18Н10Т", "ГОСТ 5632"])
    wb.save(path)
    parsed = ParserRouter().parse_document(str(path), doc_id="xlsx")
    text = _text(parsed.chunks)
    assert "material: 12Х18Н10Т" in text
    assert "standard: ГОСТ 5632" in text


def test_docx_parser_extracts_text_and_table(tmp_path: Path) -> None:
    import docx

    path = tmp_path / "pump.docx"
    doc = docx.Document()
    doc.add_paragraph("Насос NPK-200 соответствует ISO 9001.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "part"
    table.cell(0, 1).text = "article_number"
    table.cell(1, 0).text = "корпус"
    table.cell(1, 1).text = "ART-NPK-200-BODY"
    doc.save(path)
    parsed = ParserRouter().parse_document(str(path), doc_id="docx")
    text = _text(parsed.chunks)
    assert "Насос NPK-200" in text
    assert "article_number: ART-NPK-200-BODY" in text


def test_html_parser_extracts_images_and_table(tmp_path: Path) -> None:
    path = tmp_path / "valve.html"
    path.write_text(
        '<html><body><h1>Клапан DN50</h1><figure><img src="scheme.png" alt="схема монтажа"><figcaption>Схема монтажа</figcaption></figure>'
        "<table><tr><th>object</th><th>parameter</th></tr><tr><td>клапан DN50</td><td>PN16</td></tr></table></body></html>",
        encoding="utf-8",
    )
    parsed = ParserRouter().parse_document(str(path), doc_id="html")
    text = _text(parsed.chunks)
    assert "Image: url: scheme.png" in text
    assert "object: клапан DN50" in text
    assert parsed.metadata["image_refs"][0]["alt"] == "схема монтажа"
